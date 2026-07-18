# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OVERTIME_TEMPLATE = "加班补录空白模板.docx"
PENALTY_TEMPLATE = "罚班记录空白模板.docx"
OVERTIME_HEADERS = ["序号", "姓名", "加班班次"]
PENALTY_HEADERS = ["序号", "姓名", "罚班次数"]
BLANK_ROWS = 20


def set_cell_text(cell, text):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "宋体"
            run.font.size = Pt(11)


def create_adjustment_template(output_path, title, headers):
    doc = Document()
    title_paragraph = doc.add_paragraph(title)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.runs[0]
    run.font.name = "宋体"
    run.font.size = Pt(16)
    run.font.bold = True

    table = doc.add_table(rows=1 + BLANK_ROWS, cols=len(headers))
    table.style = "Table Grid"
    for col_idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[col_idx], header)

    for row_idx in range(1, BLANK_ROWS + 1):
        set_cell_text(table.rows[row_idx].cells[0], str(row_idx))
        set_cell_text(table.rows[row_idx].cells[1], "")
        set_cell_text(table.rows[row_idx].cells[2], "")

    doc.save(output_path)


def create_default_templates(output_dir=".", overwrite=False):
    from pathlib import Path

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    templates = [
        (output_dir / OVERTIME_TEMPLATE, "月统计加班补录", OVERTIME_HEADERS),
        (output_dir / PENALTY_TEMPLATE, "月统计罚班记录", PENALTY_HEADERS),
    ]
    for output_path, title, headers in templates:
        if output_path.exists() and not overwrite:
            continue
        create_adjustment_template(output_path, title, headers)
    return [path for path, _title, _headers in templates]


def main():
    for output_path in create_default_templates():
        print(output_path)


if __name__ == "__main__":
    main()
