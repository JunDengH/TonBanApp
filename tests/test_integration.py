import pytest
import pandas as pd
from docx import Document
from src.modules.data_manager import DataManager
from src.modules.word_parser import parse_total_name_list, parse_previous_month_word
from src.modules.excel_generator import (
    generate_weekly_excel,
    generate_monthly_excel,
)
from src.modules.word_generator import generate_monthly_word


def test_end_to_end(make_total_names_docx, make_weekly_docx, tmp_path):
    # 1. 导入总名单
    total = ["张三", "李四", "王五", "赵六"]
    total_path = make_total_names_docx(total)
    dm = DataManager(tmp_path / "config.json")
    dm.update_name_list(parse_total_name_list(total_path))
    assert dm.get_total_count() == 4

    # 2. 生成4个周统计
    week_files = []
    for wk in range(4):
        schedule = {
            "周一": [["张三", "李四"]],
            "周二": [["王五", "赵六"]],
        }
        word_path = make_weekly_docx(schedule)
        # 重命名避免覆盖
        import shutil
        renamed = tmp_path / f"weekly_{wk}.docx"
        shutil.copy(word_path, renamed)

        out = tmp_path / f"week{wk+1}.xlsx"
        generate_weekly_excel(
            word_path=str(renamed),
            total_names=dm.get_name_list(),
            holidays=[],
            output_path=str(out),
            actual_word_path=str(renamed),
        )
        week_files.append(str(out))

    # 3. 生成月统计
    month_out = tmp_path / "month.xlsx"
    generate_monthly_excel(
        weekly_excel_paths=week_files,
        total_names=dm.get_name_list(),
        output_path=str(month_out),
    )

    df = pd.read_excel(month_out)
    assert len(df) == 4
    assert set(df["姓名"]) == set(total)
    # 每人每周实际1次 × 4周 = 4；应值按排班统计也为4
    assert (df["总计班次"] == 4).all()
    assert (df["应值班次"] == 4).all()


def test_end_to_end_word_with_carryover(
    make_total_names_docx, make_weekly_docx, tmp_path
):
    """
    需求3 端到端：
    1) 导入总名单
    2) 生成 4 份周 Excel
    3) 首次生成月 Word（无上月）
    4) 把刚生成的 Word 作为"上月"再次生成下一月 Word
    5) 验证数据正确继承
    """
    import shutil
    total = ["张三", "李四", "王五", "赵六"]
    total_path = make_total_names_docx(total)
    dm = DataManager(tmp_path / "config.json")
    dm.update_name_list(parse_total_name_list(total_path))

    # 4 份周 Excel（每人每周 1 次）
    week_files = []
    for wk in range(4):
        schedule = {
            "周一": [["张三", "李四"]],
            "周二": [["王五", "赵六"]],
        }
        word_path = make_weekly_docx(schedule)
        renamed = tmp_path / f"weekly_{wk}.docx"
        shutil.copy(word_path, renamed)
        out = tmp_path / f"week{wk+1}.xlsx"
        generate_weekly_excel(
            word_path=str(renamed),
            total_names=dm.get_name_list(),
            holidays=[],
            output_path=str(out),
            actual_word_path=str(renamed),
        )
        week_files.append(str(out))

    # 首月 Word（无上月）
    first_month = tmp_path / "month1.docx"
    generate_monthly_word(week_files, dm.get_name_list(), None, str(first_month))
    parsed1 = parse_previous_month_word(str(first_month))
    # 首月每人：实际4 应值4 -> 多值0/缺班0
    for name in total:
        assert parsed1[name]["多值班次"] == 0
        assert parsed1[name]["缺班数量"] == 0

    # 次月 Word（继承首月缺班4）— 再来一轮同样排班
    second_month = tmp_path / "month2.docx"
    generate_monthly_word(week_files, dm.get_name_list(),
                         str(first_month), str(second_month))
    doc = Document(str(second_month))
    t = doc.tables[0]
    by = {row.cells[1].text: row for row in t.rows[2:]}
    # 每人：本月应值4 + 上月缺班0 = 4；本月实际4 + 上月多值0 = 4；缺班 = 0
    for name in total:
        assert by[name].cells[2].text == "4"   # 总计
        assert by[name].cells[3].text == "4"   # 应值
        assert by[name].cells[4].text.strip() == ""  # 多值 <=0 斜线
        assert by[name].cells[5].text.strip() == ""  # 缺班 <=0 斜线


def test_monthly_overtime_end_to_end(make_total_names_docx, make_weekly_docx, tmp_path):
    """端到端：月统计加班补录会影响最终总计班次"""
    total = ["张三", "李四"]
    total_path = make_total_names_docx(total)
    dm = DataManager(tmp_path / "config.json")
    dm.update_name_list(parse_total_name_list(total_path))

    schedule = {
        "周一": [["张三"]],  # 原始：张三1，李四0
    }
    word_path = make_weekly_docx(schedule)
    out = tmp_path / "week.xlsx"

    generate_weekly_excel(
        word_path=str(word_path),
        total_names=dm.get_name_list(),
        holidays=[],
        output_path=str(out),
        actual_word_path=str(word_path),
    )

    month_out = tmp_path / "month_with_overtime.docx"
    generate_monthly_word(
        weekly_excel_paths=[str(out)],
        total_names=dm.get_name_list(),
        prev_word_path=None,
        output_path=str(month_out),
        overtime_shifts={"张三": 2, "李四": 1},
    )

    doc = Document(str(month_out))
    t = doc.tables[0]
    by = {row.cells[1].text: row for row in t.rows[2:]}
    # 原本：张三实际1、李四实际0；月加班后总计应分别 +2 / +1
    assert by["张三"].cells[2].text == "3"
    assert by["李四"].cells[2].text == "1"