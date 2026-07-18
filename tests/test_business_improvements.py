import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import load_workbook, Workbook

from src.modules.contact_parser import compute_graduation_seniors, parse_contact_xlsx
from src.modules.data_manager import DataManager
from src.modules.excel_generator import (
    build_weekly_rows,
    build_weekly_rows_from_schedules,
    collect_typo_suspects,
    collect_weekly_warnings,
    generate_weekly_excel_from_rows,
)
from src.modules.word_generator import build_monthly_rows, collect_monthly_warnings
from src.modules.word_parser import (
    find_typo_suspects,
    parse_previous_month_word,
    scan_weekly_word,
)
from src.utils.helpers import chars_sound_similar, edit_distance, pinyin_readings
from src.utils.output_paths import build_unique_output_path
from src.utils.report_titles import (
    build_final_period_report_title,
    build_monthly_report_title,
    build_weekly_report_title,
)


def make_docx(path: Path, lines: list[str]) -> None:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(path)


def make_weekly_xlsx(path: Path, rows: list[tuple[str, int, int]]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "周统计"
    ws.append(["姓名", "应值班次", "实际班次", "缺班", "备注"])
    for name, should, actual in rows:
        ws.append([name, should, actual, max(0, should - actual), ""])
    wb.save(path)


def make_contact_xlsx(path: Path) -> None:
    """构造一个含合并风格部门、全角空格姓名、反序/带空格专业年级的通讯录。"""
    wb = Workbook()
    ws = wb.active
    ws.cell(row=1, column=1, value="2026年通讯录")  # 标题行
    for col, header in enumerate(["部门", "姓名", "专业年级", "联系方式"], start=1):
        ws.cell(row=2, column=col, value=header)
    rows = [
        ("阳光俱乐部", "张　爽", "新闻22级"),   # 部门块首行 + 全角空格姓名
        (None, "周　密", "社会22级"),           # 部门需向下填充
        (None, "王海韵", "24级法学"),           # 反序
        ("义工部", "李四", "材料 24级"),         # 新部门块 + 带空格
        (None, "钱五", "管理23级"),
    ]
    for r, (dept, name, pg) in enumerate(rows, start=3):
        if dept is not None:
            ws.cell(row=r, column=1, value=dept)
        ws.cell(row=r, column=2, value=name)
        ws.cell(row=r, column=3, value=pg)
    wb.save(path)


class BusinessImprovementTests(unittest.TestCase):
    def test_build_weekly_rows_from_schedules_applies_final_week_multiplier(self):
        rows = build_weekly_rows_from_schedules(
            schedule_for_should={"周一": ["张三"]},
            schedule_for_actual={"周一": ["张三"]},
            total_names=["张三"],
            holidays=[],
            senior_assistants=[],
            senior_should_fixed_enabled="normal",
            long_term_leave_assistants=[],
            final_week_enabled=True,
        )

        self.assertEqual(rows, [{
            "姓名": "张三",
            "应值班次": 2,
            "实际班次": 2,
            "缺班": 0,
            "备注": "",
        }])

    def test_weekly_tab_never_reads_final_week_switch(self):
        source = Path("src/ui/weekly_tab.py").read_text(encoding="utf-8")

        self.assertNotIn("get_final_week_enabled", source)
        self.assertIn("final_week_enabled=False", source)

    def test_data_manager_persists_theme_and_output_dir(self):
        with tempfile.TemporaryDirectory() as tmp:
            config_path = Path(tmp) / "data" / "config.json"
            output_dir = Path(tmp) / "exports"

            manager = DataManager(config_path=config_path)
            self.assertEqual(manager.get_app_theme(), "dark")
            self.assertEqual(Path(manager.get_output_dir()).name, "output")

            manager.set_app_theme("light")
            manager.set_output_dir(str(output_dir))

            reloaded = DataManager(config_path=config_path)
            self.assertEqual(reloaded.get_app_theme(), "light")
            self.assertEqual(Path(reloaded.get_output_dir()), output_dir)

    def test_name_majors_persist_and_clean_with_name_list(self):
        with tempfile.TemporaryDirectory() as tmp:
            config_path = Path(tmp) / "data" / "config.json"
            manager = DataManager(config_path=config_path)
            self.assertEqual(manager.get_name_majors(), {})

            manager.update_name_list(["张爽", "周密"])
            manager._data["name_majors"] = {"张爽": "新闻", "周密": "社会"}
            manager._data["name_grades"] = {"张爽": "22级", "周密": "22级"}
            manager._save()

            reloaded = DataManager(config_path=config_path)
            self.assertEqual(reloaded.get_name_major("张爽"), "新闻")
            self.assertEqual(reloaded.get_name_grade("周密"), "22级")

            # 周密 离开名单后，其专业/年级应被清理；张爽保留
            reloaded.update_name_list(["张爽"])
            self.assertEqual(reloaded.get_name_majors(), {"张爽": "新闻"})
            self.assertEqual(reloaded.get_name_grade("周密"), "")

    def test_parse_contact_xlsx_recognizes_all_fields(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "contact.xlsx"
            make_contact_xlsx(path)
            records = parse_contact_xlsx(str(path))
            by_name = {r["name"]: r for r in records}
            self.assertEqual(len(records), 5)
            # 全角空格姓名归一
            self.assertIn("张爽", by_name)
            self.assertEqual(by_name["张爽"]["major"], "新闻")
            self.assertEqual(by_name["张爽"]["grade"], "22级")
            # 部门向下填充
            self.assertEqual(by_name["周密"]["department"], "阳光俱乐部")
            self.assertEqual(by_name["钱五"]["department"], "义工部")
            # 反序 专业年级
            self.assertEqual(by_name["王海韵"]["major"], "法学")
            self.assertEqual(by_name["王海韵"]["grade"], "24级")
            # 带空格 专业年级
            self.assertEqual(by_name["李四"]["major"], "材料")
            self.assertEqual(by_name["李四"]["grade"], "24级")

    def test_compute_graduation_seniors_picks_earliest_cohort(self):
        records = [
            {"name": "甲", "grade": "22级"},
            {"name": "乙", "grade": "22级"},
            {"name": "丙", "grade": "23级"},
            {"name": "丁", "grade": "24级"},  # 共 3 种年级，仍应识别
        ]
        self.assertEqual(compute_graduation_seniors(records), ["甲", "乙"])
        self.assertEqual(compute_graduation_seniors([]), [])

    def test_apply_contact_roster_sets_fields_and_overrides_seniors(self):
        with tempfile.TemporaryDirectory() as tmp:
            manager = DataManager(config_path=Path(tmp) / "data" / "config.json")
            # 预置：旧名单 + 旧毕业季 + 长期请假（验证覆盖与按名保留）
            manager.update_name_list(["张爽", "某离职"])
            manager.set_name_senior("某离职", True)
            manager.set_name_long_term_leave("张爽", True)

            records = [
                {"name": "张爽", "department": "阳光俱乐部", "major": "新闻", "grade": "22级"},
                {"name": "李四", "department": "义工部", "major": "材料", "grade": "24级"},
            ]
            seniors = compute_graduation_seniors(records)  # ["张爽"]
            manager.apply_contact_roster(records, seniors)

            self.assertEqual(manager.get_name_list(), ["张爽", "李四"])
            self.assertEqual(manager.get_name_department("张爽"), "阳光俱乐部")
            self.assertEqual(manager.get_name_major("李四"), "材料")
            self.assertEqual(manager.get_name_grade("李四"), "24级")
            # 毕业季被覆盖为最靠前年级；旧毕业季“某离职”不再存在
            self.assertEqual(manager.get_senior_assistants(), ["张爽"])
            # 长期请假按姓名保留（张爽仍在新名单）
            self.assertEqual(manager.get_long_term_leave_assistants(), ["张爽"])

    def test_name_departments_persist_and_clean_with_name_list(self):
        with tempfile.TemporaryDirectory() as tmp:
            config_path = Path(tmp) / "data" / "config.json"
            manager = DataManager(config_path=config_path)
            self.assertEqual(manager.get_name_departments(), {})
            manager.update_name_list(["甲", "乙"])
            manager.set_name_department("甲", "行政部")
            manager.set_name_department("乙", "新闻部")
            self.assertEqual(DataManager(config_path=config_path).get_name_department("甲"), "行政部")
            # 乙 离开名单后其部门应被清理
            manager.update_name_list(["甲"])
            self.assertEqual(manager.get_name_departments(), {"甲": "行政部"})

    def test_edit_rename_remove_and_rules(self):
        with tempfile.TemporaryDirectory() as tmp:
            manager = DataManager(config_path=Path(tmp) / "data" / "config.json")
            manager.update_name_list(["张三", "李四"])
            manager.add_name("王五", grade="25级", major="计科")
            self.assertEqual(manager.get_name_major("王五"), "计科")

            manager.set_name_senior("张三", True)
            manager.set_name_long_term_leave("张三", True)
            # 重命名应迁移全部引用并保持位置
            self.assertTrue(manager.rename_name("张三", "张三丰"))
            self.assertEqual(manager.get_name_list()[0], "张三丰")
            self.assertIn("张三丰", manager.get_senior_assistants())
            self.assertIn("张三丰", manager.get_long_term_leave_assistants())
            # 重名应被拒绝
            self.assertFalse(manager.rename_name("张三丰", "李四"))
            # 移出毕业季
            manager.set_name_senior("张三丰", False)
            self.assertNotIn("张三丰", manager.get_senior_assistants())
            # 删除应级联清理
            self.assertTrue(manager.remove_name("李四"))
            self.assertNotIn("李四", manager.get_name_list())

    def test_special_recognize_removes_all_seniors(self):
        with tempfile.TemporaryDirectory() as tmp:
            manager = DataManager(config_path=Path(tmp) / "data" / "config.json")
            manager.update_name_list(["甲一", "乙二", "丙三"])
            manager.set_name_senior("甲一", True)
            manager.set_name_senior("丙三", True)
            manager.set_name_long_term_leave("甲一", True)

            seniors = manager.get_senior_assistants()
            remaining = [n for n in manager.get_name_list() if n not in set(seniors)]
            manager.update_name_list(remaining)

            self.assertEqual(manager.get_name_list(), ["乙二"])
            self.assertEqual(manager.get_senior_assistants(), [])
            self.assertEqual(manager.get_long_term_leave_assistants(), [])

    def test_unique_output_path_appends_number_without_overwriting(self):
        with tempfile.TemporaryDirectory() as tmp:
            output_dir = Path(tmp)
            first = output_dir / "2026春季学期第1周助理值班统计.xlsx"
            first.write_text("existing", encoding="utf-8")

            result = build_unique_output_path(
                output_dir,
                "2026春季学期第1周助理值班统计",
                ".xlsx",
            )

            self.assertEqual(result.name, "2026春季学期第1周助理值班统计 (2).xlsx")
            self.assertEqual(first.read_text(encoding="utf-8"), "existing")

    def test_default_report_filenames_follow_title_rules(self):
        with tempfile.TemporaryDirectory() as tmp:
            output_dir = Path(tmp)
            weekly_title = build_weekly_report_title("2026", "春季", "1")
            monthly_title = build_monthly_report_title("2026", "秋季", "5", "8")

            weekly_path = build_unique_output_path(output_dir, weekly_title, ".xlsx")
            monthly_path = build_unique_output_path(output_dir, monthly_title, ".docx")

            self.assertEqual(weekly_path.name, "2026春季学期第1周助理值班统计.xlsx")
            self.assertEqual(monthly_path.name, "2026秋季学期5-8周助理值班统计.docx")

    def test_final_period_title_has_no_week_number(self):
        self.assertEqual(
            build_final_period_report_title("2026", "春季"),
            "2026春季学期期末周助理值班统计",
        )

    def test_build_weekly_rows_and_generate_excel_from_preview_rows(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            schedule_path = tmp_path / "schedule.docx"
            actual_path = tmp_path / "actual.docx"
            output_path = tmp_path / "weekly.xlsx"
            make_docx(schedule_path, ["周一", "张三 李四", "周二", "张三"])
            make_docx(actual_path, ["周一", "张三"])

            rows = build_weekly_rows(
                word_path=str(schedule_path),
                total_names=["张三", "李四"],
                holidays=[],
                actual_word_path=str(actual_path),
                senior_assistants=[],
                senior_should_fixed_enabled="normal",
                long_term_leave_assistants=[],
            )
            self.assertEqual(rows[0]["姓名"], "张三")
            self.assertEqual(rows[0]["应值班次"], 2)
            self.assertEqual(rows[0]["实际班次"], 1)
            self.assertEqual(rows[0]["缺班"], 1)

            rows[0]["实际班次"] = 2
            rows[0]["备注"] = "人工核对"
            generate_weekly_excel_from_rows(rows, str(output_path))

            wb = load_workbook(output_path)
            ws = wb["周统计"]
            self.assertEqual(ws["A1"].value, "姓名")
            excel_rows = {
                ws.cell(row=row, column=1).value: {
                    "实际班次": ws.cell(row=row, column=3).value,
                    "缺班": ws.cell(row=row, column=4).value,
                    "备注": ws.cell(row=row, column=5).value,
                }
                for row in range(2, ws.max_row + 1)
            }
            self.assertEqual(excel_rows["张三"]["实际班次"], 2)
            self.assertEqual(excel_rows["张三"]["缺班"], 0)
            self.assertEqual(excel_rows["张三"]["备注"], "人工核对")

    def test_final_week_rule_doubles_should_and_actual(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            schedule_path = tmp_path / "schedule.docx"
            actual_path = tmp_path / "actual.docx"
            # 期末周：每人在表上只出现一次，但李四在值班表上出现两次（超额场景）。
            make_docx(schedule_path, ["周一", "张三 李四"])
            make_docx(actual_path, ["周一", "张三 李四", "周二", "李四"])

            rows = build_weekly_rows(
                word_path=str(schedule_path),
                total_names=["张三", "李四", "王五"],
                holidays=[],
                actual_word_path=str(actual_path),
                senior_assistants=[],
                senior_should_fixed_enabled="normal",
                long_term_leave_assistants=[],
                final_week_enabled=True,
            )
            by_name = {r["姓名"]: r for r in rows}
            # 张三：排版1次→应值2，值班1次→实际2，缺班0
            self.assertEqual(by_name["张三"]["应值班次"], 2)
            self.assertEqual(by_name["张三"]["实际班次"], 2)
            self.assertEqual(by_name["张三"]["缺班"], 0)
            # 李四：排版1次→应值2，值班2次→实际4，缺班0（多值）
            self.assertEqual(by_name["李四"]["应值班次"], 2)
            self.assertEqual(by_name["李四"]["实际班次"], 4)
            self.assertEqual(by_name["李四"]["缺班"], 0)
            # 王五：未排版未值班 → 全 0
            self.assertEqual(by_name["王五"]["应值班次"], 0)
            self.assertEqual(by_name["王五"]["实际班次"], 0)

    def test_final_week_holiday_reduction_doubles(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            schedule_path = tmp_path / "schedule.docx"
            actual_path = tmp_path / "actual.docx"
            make_docx(schedule_path, ["周一", "张三"])
            make_docx(actual_path, [])

            rows = build_weekly_rows(
                word_path=str(schedule_path),
                total_names=["张三"],
                holidays=["周一"],
                actual_word_path=str(actual_path),
                senior_assistants=[],
                senior_should_fixed_enabled="normal",
                long_term_leave_assistants=[],
                final_week_enabled=True,
            )
            # 排版1次→应值2，周一放假核减按双倍 → max(0, 2 - 1*2) = 0
            self.assertEqual(rows[0]["姓名"], "张三")
            self.assertEqual(rows[0]["应值班次"], 0)

    def test_final_week_fixed_quota_not_doubled_but_actual_doubled(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            schedule_path = tmp_path / "schedule.docx"
            actual_path = tmp_path / "actual.docx"
            make_docx(schedule_path, [])
            # 长期请假者赵六到岗1次；毕业季无需值班者孙七到岗1次。
            make_docx(actual_path, ["周一", "赵六 孙七"])

            rows = build_weekly_rows(
                word_path=str(schedule_path),
                total_names=["赵六", "孙七"],
                holidays=[],
                actual_word_path=str(actual_path),
                senior_assistants=["孙七"],
                senior_should_fixed_enabled="none",
                long_term_leave_assistants=["赵六"],
                final_week_enabled=True,
            )
            by_name = {r["姓名"]: r for r in rows}
            # 长期请假：应值固定2（不翻倍），实际1次→2（翻倍），缺班0
            self.assertEqual(by_name["赵六"]["应值班次"], 2)
            self.assertEqual(by_name["赵六"]["实际班次"], 2)
            self.assertEqual(by_name["赵六"]["缺班"], 0)
            # 毕业季无需值班：应值0，实际1次→2（翻倍）
            self.assertEqual(by_name["孙七"]["应值班次"], 0)
            self.assertEqual(by_name["孙七"]["实际班次"], 2)

    def test_final_week_conflicts_with_reduced_senior_mode(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            schedule_path = tmp_path / "schedule.docx"
            actual_path = tmp_path / "actual.docx"
            make_docx(schedule_path, ["周一", "周八"])
            make_docx(actual_path, ["周一", "周八"])

            with self.assertRaises(ValueError):
                build_weekly_rows(
                    word_path=str(schedule_path),
                    total_names=["周八"],
                    holidays=[],
                    actual_word_path=str(actual_path),
                    senior_assistants=["周八"],
                    senior_should_fixed_enabled="reduced",
                    long_term_leave_assistants=[],
                    final_week_enabled=True,
                )

    def test_final_week_enabled_persists_and_defaults_false(self):
        with tempfile.TemporaryDirectory() as tmp:
            config_path = Path(tmp) / "config.json"
            manager = DataManager(config_path=config_path)
            # 默认 False
            self.assertFalse(manager.get_final_week_enabled())
            manager.set_final_week_enabled(True)
            # 重新加载仍为 True
            reloaded = DataManager(config_path=config_path)
            self.assertTrue(reloaded.get_final_week_enabled())

    def test_scan_weekly_word_counts_matched_and_unknown(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            doc_path = tmp_path / "schedule.docx"
            make_docx(doc_path, ["周一", "张三 王五", "周二", "陌生人"])

            result = scan_weekly_word(str(doc_path), ["张三"])
            self.assertEqual(result["matched_count"], 1)
            self.assertIn("王五", result["unknown_names"])
            self.assertIn("陌生人", result["unknown_names"])
            self.assertNotIn("张三", result["unknown_names"])

    def test_scan_weekly_word_detects_no_match(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            doc_path = tmp_path / "wrong.docx"
            make_docx(doc_path, ["周一", "陌生人 路人甲"])

            result = scan_weekly_word(str(doc_path), ["张三", "李四"])
            self.assertEqual(result["matched_count"], 0)

    def test_collect_weekly_warnings_flags_should_over_two(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            schedule_path = tmp_path / "schedule.docx"
            actual_path = tmp_path / "actual.docx"
            make_docx(schedule_path, ["周一", "张三 李四"])
            make_docx(actual_path, ["周一", "张三 李四"])

            rows = [
                {"姓名": "张三", "应值班次": 3, "实际班次": 1, "缺班": 2, "备注": ""},
                {"姓名": "李四", "应值班次": 2, "实际班次": 2, "缺班": 0, "备注": ""},
            ]
            warnings = collect_weekly_warnings(rows, str(schedule_path), str(actual_path), ["张三", "李四"])
            self.assertIn("张三", warnings["highlight_names"])
            self.assertNotIn("李四", warnings["highlight_names"])
            self.assertTrue(warnings["messages"])

    def test_pinyin_readings_handles_heteronym(self):
        # 多音字“乐”应同时给出 le / yue 等读音。
        readings = pinyin_readings("乐")
        self.assertIn("le", readings)
        self.assertIn("yue", readings)

    def test_chars_sound_similar_fuzzy_and_distinct(self):
        # 模糊音命中：lin/ling、zhou/zou（差 1 个字母）。
        self.assertTrue(chars_sound_similar("林", "玲"))
        self.assertTrue(chars_sound_similar("周", "邹"))
        # 同音也算相近。
        self.assertTrue(chars_sound_similar("爽", "霜"))
        # 读音明显不同：ming vs gang 不相近。
        self.assertFalse(chars_sound_similar("明", "刚"))

    def test_edit_distance_basic(self):
        self.assertEqual(edit_distance("王晓明", "王明"), 1)   # 漏字
        self.assertEqual(edit_distance("张三", "李四"), 2)

    def test_find_typo_suspects_strong_homophone(self):
        # 整名同音、用字不同 -> 高度疑似（声调不同仍同音）。
        suspects = find_typo_suspects(["张霜", "李四"], ["张爽"])
        result = {s["name"]: s for s in suspects}
        self.assertIn("张霜", result)
        self.assertEqual(result["张霜"]["level"], "strong")
        self.assertEqual(result["张霜"]["candidates"], ["张爽"])
        self.assertNotIn("李四", result)         # 与“张爽”读音/字形都不沾边

    def test_find_typo_suspects_medium_missing_or_extra_char(self):
        # 漏字 / 多字 -> 疑似。
        suspects = {s["name"]: s["level"] for s in find_typo_suspects(["王明", "王晓晓明"], ["王晓明"])}
        self.assertEqual(suspects.get("王明"), "medium")
        self.assertEqual(suspects.get("王晓晓明"), "medium")

    def test_find_typo_suspects_medium_fuzzy_pinyin(self):
        # 替换且读音相近（in/ing）-> 疑似。
        suspects = {s["name"]: s["level"] for s in find_typo_suspects(["李玲"], ["李林"])}
        self.assertEqual(suspects.get("李玲"), "medium")

    def test_find_typo_suspects_weak_shape_only_three_char(self):
        # 三字：替换但读音不同 -> 可能（weak）。wei 与 ming/meng 读音都对不上。
        suspects = {s["name"]: s["level"] for s in find_typo_suspects(["王晓伟"], ["王晓明"])}
        self.assertEqual(suspects.get("王晓伟"), "weak")

    def test_find_typo_suspects_two_char_shape_only_reports_weak(self):
        # 两字：替换且读音不同 -> 也报（weak，避免遗漏）；读音相近/同音照常。
        suspects = {s["name"]: s["level"] for s in find_typo_suspects(["王鹏", "王方"], ["王芳"])}
        self.assertEqual(suspects.get("王鹏"), "weak")    # peng/fang 读音不同 -> 形近 weak
        self.assertEqual(suspects.get("王方"), "strong")  # fang/fang 同音

    def test_collect_typo_suspects_scans_both_sources(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            schedule_path = tmp_path / "schedule.docx"
            actual_path = tmp_path / "actual.docx"
            make_docx(schedule_path, ["周一", "张霜 李四"])      # 张爽 -> 张霜（strong）
            make_docx(actual_path, ["周一", "张爽 王晓伟"])       # 王晓明 -> 王晓伟（weak）

            suspects = collect_typo_suspects(str(schedule_path), str(actual_path), ["张爽", "李四", "王晓明"])
            by_name = {s["name"]: s for s in suspects}
            self.assertEqual(by_name["张霜"]["source"], "排班 Word")
            self.assertEqual(by_name["张霜"]["level"], "strong")
            self.assertEqual(by_name["王晓伟"]["source"], "实际 Word")
            self.assertEqual(by_name["王晓伟"]["level"], "weak")

    def test_collect_typo_suspects_skips_missing_path(self):
        # 路径为空或不存在时静默跳过，不抛异常。
        self.assertEqual(collect_typo_suspects("", None, ["张爽"]), [])

    def test_collect_typo_suspects_dedupes_same_file(self):
        # 同一文件被同时选为排班与实际时只算一次，不重复计数。
        with tempfile.TemporaryDirectory() as tmp:
            doc_path = Path(tmp) / "same.docx"
            make_docx(doc_path, ["周一", "张霜 李四"])
            suspects = collect_typo_suspects(str(doc_path), str(doc_path), ["张爽", "李四"])
            self.assertEqual(len(suspects), 1)
            self.assertEqual(suspects[0]["source"], "排班 Word")

    def test_collect_weekly_warnings_reports_typo_separately(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            schedule_path = tmp_path / "schedule.docx"
            actual_path = tmp_path / "actual.docx"
            # 排班表把“张爽”打成了“张霜”，并夹带一段说明文字。
            make_docx(schedule_path, ["周一", "张霜 李四 值班说明"])
            make_docx(actual_path, ["周一", "张爽 李四"])

            rows = [
                {"姓名": "张爽", "应值班次": 1, "实际班次": 1, "缺班": 0, "备注": ""},
                {"姓名": "李四", "应值班次": 1, "实际班次": 1, "缺班": 0, "备注": ""},
            ]
            warnings = collect_weekly_warnings(
                rows, str(schedule_path), str(actual_path), ["张爽", "李四"]
            )
            # “张霜”归入高度疑似/疑似通道，并带 level。
            self.assertTrue(warnings["typo_messages_strong"])
            self.assertEqual(warnings["typo_suspects"][0]["name"], "张霜")
            self.assertEqual(warnings["typo_suspects"][0]["level"], "strong")
            self.assertIn("张爽", warnings["typo_suspects"][0]["candidates"])
            # 已被识别为疑似错字的“张霜”，不再出现在通用“已忽略”提示里。
            self.assertNotIn("张霜", "".join(warnings["messages"]))
            # “值班说明”不是疑似错字，仍按“已忽略”提示保留。
            self.assertIn("值班说明", "".join(warnings["messages"]))

    def test_collect_weekly_warnings_clean_case_has_no_messages(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            schedule_path = tmp_path / "schedule.docx"
            actual_path = tmp_path / "actual.docx"
            make_docx(schedule_path, ["周一", "张三 李四"])
            make_docx(actual_path, ["周一", "张三 李四"])

            rows = [
                {"姓名": "张三", "应值班次": 2, "实际班次": 2, "缺班": 0, "备注": ""},
                {"姓名": "李四", "应值班次": 1, "实际班次": 1, "缺班": 0, "备注": ""},
            ]
            warnings = collect_weekly_warnings(rows, str(schedule_path), str(actual_path), ["张三", "李四"])
            self.assertEqual(warnings["highlight_names"], set())
            self.assertEqual(warnings["messages"], [])

    def test_collect_monthly_warnings_range_only_at_four_weeks(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            total_names = ["甲", "乙", "丙", "丁"]
            # 甲=10(>8) 乙=8(边界) 丙=4(边界) 丁=10但毕业季豁免
            week_rows = [
                [("甲", 3, 3), ("乙", 2, 2), ("丙", 1, 1), ("丁", 3, 3)],
                [("甲", 3, 3), ("乙", 2, 2), ("丙", 1, 1), ("丁", 3, 3)],
                [("甲", 2, 2), ("乙", 2, 2), ("丙", 1, 1), ("丁", 2, 2)],
                [("甲", 2, 2), ("乙", 2, 2), ("丙", 1, 1), ("丁", 2, 2)],
            ]
            paths = []
            for i, rows in enumerate(week_rows):
                p = tmp_path / f"week{i}.xlsx"
                make_weekly_xlsx(p, rows)
                paths.append(str(p))

            warnings = collect_monthly_warnings(paths, total_names, senior_assistants=["丁"])
            self.assertIn("甲", warnings["highlight_names"])  # 10 > 8
            self.assertNotIn("乙", warnings["highlight_names"])  # 边界 8
            self.assertNotIn("丙", warnings["highlight_names"])  # 边界 4
            self.assertNotIn("丁", warnings["highlight_names"])  # 毕业季豁免

    def test_collect_monthly_warnings_skips_range_when_not_four_weeks(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            total_names = ["甲", "乙", "丙"]
            # 仅 2 周：甲周应值合计=10，但非满 4 周不做范围校验；乙全零→命中；丙毕业季全零→豁免
            week_rows = [
                [("甲", 5, 5), ("乙", 0, 0), ("丙", 0, 0)],
                [("甲", 5, 5), ("乙", 0, 0), ("丙", 0, 0)],
            ]
            paths = []
            for i, rows in enumerate(week_rows):
                p = tmp_path / f"week{i}.xlsx"
                make_weekly_xlsx(p, rows)
                paths.append(str(p))

            warnings = collect_monthly_warnings(paths, total_names, senior_assistants=["丙"])
            self.assertNotIn("甲", warnings["highlight_names"])  # 仅2周，不做范围校验
            self.assertIn("乙", warnings["highlight_names"])  # 全零命中
            self.assertNotIn("丙", warnings["highlight_names"])  # 毕业季豁免

    def test_monthly_rows_accepts_one_to_four_weekly_files(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            files = []
            for index in range(4):
                path = tmp_path / f"week-{index + 1}.xlsx"
                make_weekly_xlsx(path, [("张三", 2, 1)])
                files.append(str(path))

            one_week = build_monthly_rows(files[:1], ["张三"], None)
            two_weeks = build_monthly_rows(files[:2], ["张三"], None)
            four_weeks = build_monthly_rows(files[:4], ["张三"], None)

            self.assertEqual(one_week[0]["应值班次"], 2)
            self.assertEqual(two_weeks[0]["应值班次"], 4)
            self.assertEqual(four_weeks[0]["应值班次"], 8)
            with self.assertRaisesRegex(ValueError, "至少选择 1 份"):
                build_monthly_rows([], ["张三"], None)
            with self.assertRaisesRegex(ValueError, "最多选择 4 份"):
                build_monthly_rows(files + [files[0]], ["张三"], None)

    def test_previous_month_negative_cells_are_clamped_to_zero(self):
        """上月 Word 被手工改成负数时，结转应按 0 处理（防篡改）。"""
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "prev.docx"
            doc = Document()
            table = doc.add_table(rows=2, cols=6)
            headers = ["序号", "姓名", "总计班次", "应值班次", "多值班次", "缺班数量"]
            for col, header in enumerate(headers):
                table.rows[0].cells[col].text = header
            for col, value in enumerate(["1", "张三", "5", "3", "-2", "-9"]):
                table.rows[1].cells[col].text = value
            doc.save(path)

            parsed = parse_previous_month_word(str(path))
            self.assertEqual(parsed["张三"]["多值班次"], 0)
            self.assertEqual(parsed["张三"]["缺班数量"], 0)

            files = [Path(tmp) / "w1.xlsx"]
            make_weekly_xlsx(files[0], [("张三", 2, 1)])
            rows = build_monthly_rows([str(files[0])], ["张三"], str(path))
            # 负数被钳为 0 后：总计=实际1+上月多值0；应值=应值2+上月缺班0
            self.assertEqual(rows[0]["总计班次"], 1)
            self.assertEqual(rows[0]["应值班次"], 2)

    # ===== 暂时纠错机制 =====

    def test_match_roster_names_with_corrections_maps_typo_to_correct(self):
        from src.modules.word_parser import _match_roster_names
        # "张灏琛" is NOT in roster, but correction maps it to "张颢琛"
        result = _match_roster_names("张灏琛 李四", ["张颢琛", "李四"], {"张灏琛": "张颢琛"})
        self.assertEqual(result, ["张颢琛", "李四"])

    def test_match_roster_names_without_corrections_skips_typo(self):
        from src.modules.word_parser import _match_roster_names
        # Without corrections, the typo is simply not matched
        result = _match_roster_names("张灏琛 李四", ["张颢琛", "李四"])
        self.assertEqual(result, ["李四"])

    def test_unknown_candidates_excludes_corrected_names(self):
        from src.modules.word_parser import _unknown_candidates
        # _unknown_candidates processes line-by-line, extracting compact Chinese text
        text = "张灏琛\n陌生人"
        without_corr = _unknown_candidates(text, ["张颢琛"])
        self.assertIn("张灏琛", without_corr)
        with_corr = _unknown_candidates(text, ["张颢琛"], {"张灏琛": "张颢琛"})
        self.assertNotIn("张灏琛", with_corr)
        self.assertIn("陌生人", with_corr)

    def test_scan_weekly_word_with_corrections_counts_correctly(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            doc_path = tmp_path / "schedule.docx"
            make_docx(doc_path, ["周一", "张灏琛 李四"])
            # "张灏琛" is NOT in roster, so without corrections matched=1 (only 李四)
            scan_plain = scan_weekly_word(str(doc_path), ["张颢琛", "李四"])
            self.assertEqual(scan_plain["matched_count"], 1)
            self.assertIn("张灏琛", scan_plain["unknown_names"])

            # With corrections, "张灏琛" → "张颢琛", matched=2, no unknown
            scan_corr = scan_weekly_word(
                str(doc_path), ["张颢琛", "李四"], {"张灏琛": "张颢琛"}
            )
            self.assertEqual(scan_corr["matched_count"], 2)
            self.assertNotIn("张灏琛", scan_corr["unknown_names"])

    def test_build_weekly_rows_with_corrections_counts_typo_as_correct(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            schedule_path = tmp_path / "schedule.docx"
            actual_path = tmp_path / "actual.docx"
            # Schedule has "张灏琛" (typo of "张颢琛"); actual has "张颢琛"
            make_docx(schedule_path, ["周一", "张灏琛 李四"])
            make_docx(actual_path, ["周一", "张颢琛"])

            # Without corrections: should=1 (李四), actual=1 (张颢琛)
            rows_plain = build_weekly_rows(
                word_path=str(schedule_path),
                total_names=["张颢琛", "李四"],
                holidays=[],
                actual_word_path=str(actual_path),
            )
            by_name = {r["姓名"]: r for r in rows_plain}
            self.assertEqual(by_name["李四"]["应值班次"], 1)
            self.assertEqual(by_name["张颢琛"]["应值班次"], 0)  # typo not recognized
            self.assertEqual(by_name["张颢琛"]["实际班次"], 1)

            # With corrections: should=2 (张颢琛+李四), actual=1 (张颢琛)
            rows_corr = build_weekly_rows(
                word_path=str(schedule_path),
                total_names=["张颢琛", "李四"],
                holidays=[],
                actual_word_path=str(actual_path),
                corrections={"张灏琛": "张颢琛"},
            )
            by_name = {r["姓名"]: r for r in rows_corr}
            self.assertEqual(by_name["张颢琛"]["应值班次"], 1)  # typo corrected
            self.assertEqual(by_name["李四"]["应值班次"], 1)
            self.assertEqual(by_name["张颢琛"]["实际班次"], 1)

    def test_collect_weekly_warnings_suppresses_corrected_typos(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            schedule_path = tmp_path / "schedule.docx"
            actual_path = tmp_path / "actual.docx"
            make_docx(schedule_path, ["周一", "张灏琛 李四"])
            make_docx(actual_path, ["周一", "张颢琛 李四"])

            rows = [
                {"姓名": "张颢琛", "应值班次": 1, "实际班次": 1, "缺班": 0, "备注": ""},
                {"姓名": "李四", "应值班次": 1, "实际班次": 1, "缺班": 0, "备注": ""},
            ]
            # Without corrections: "张灏琛" appears as a typo suspect
            warnings_plain = collect_weekly_warnings(
                rows, str(schedule_path), str(actual_path), ["张颢琛", "李四"]
            )
            self.assertTrue(warnings_plain["typo_messages_strong"])

            # With corrections: "张灏琛" is corrected, no typo warnings
            warnings_corr = collect_weekly_warnings(
                rows, str(schedule_path), str(actual_path), ["张颢琛", "李四"],
                corrections={"张灏琛": "张颢琛"},
            )
            self.assertEqual(warnings_corr["typo_messages_strong"], [])
            self.assertEqual(warnings_corr["typo_suspects"], [])


if __name__ == "__main__":
    unittest.main()
