import pytest
from src.modules.word_parser import (
    parse_total_name_list,
    parse_weekly_schedule,
    count_actual_shifts,
)


def test_parse_total_names(make_total_names_docx):
    names = ["张三", "李四", "王小明", "欧阳娜娜"]
    path = make_total_names_docx(names)
    result = parse_total_name_list(path)
    assert set(result) == set(names)
    assert len(result) == len(names)


def test_parse_total_names_dedup(make_total_names_docx):
    path = make_total_names_docx(["张三", "李四", "张三"])
    result = parse_total_name_list(path)
    assert result == ["张三", "李四"]


def test_parse_total_names_with_fullwidth_space(tmp_path):
    """姓名中间为全角空格（如 张　爽）时也应被识别为完整姓名"""
    from docx import Document

    path = tmp_path / "total_fullwidth.docx"
    doc = Document()
    doc.add_paragraph("张　爽 周　密 黄雪亭")
    doc.save(path)

    result = parse_total_name_list(str(path))
    assert result == ["张爽", "周密", "黄雪亭"]


def test_parse_weekly_schedule_basic(make_weekly_docx):
    schedule = {
        "周一": [["张三", "李四"], ["王五"]],
        "周二": [["李四", "王五"]],
    }
    path = make_weekly_docx(schedule)
    total = ["张三", "李四", "王五"]
    result = parse_weekly_schedule(path, total)

    assert sorted(result["周一"]) == sorted(["张三", "李四", "王五"])
    assert sorted(result["周二"]) == sorted(["李四", "王五"])
    assert result["周三"] == []


def test_parse_weekly_filters_non_names(make_weekly_docx):
    """课节号(1-2)不应进入结果"""
    schedule = {"周一": [["张三", "李四"]]}
    path = make_weekly_docx(schedule)
    result = parse_weekly_schedule(path, ["张三", "李四"])
    # 结果里只能是姓名
    for name in result["周一"]:
        assert name in {"张三", "李四"}


def test_parse_weekly_ignores_non_listed(make_weekly_docx):
    """未在总名单中的姓名应被过滤"""
    schedule = {"周一": [["张三", "路人甲"]]}
    path = make_weekly_docx(schedule)
    result = parse_weekly_schedule(path, ["张三"])
    assert result["周一"] == ["张三"]


def test_parse_weekly_header_phrase_not_name(tmp_path):
    """周标题短语（如 周一/第1周）不应被识别为姓名"""
    from docx import Document

    path = tmp_path / "weekly_header_phrase.docx"
    doc = Document()
    doc.add_paragraph("第1周")
    doc.add_paragraph("周一")
    doc.add_paragraph("1-2节 张三")
    doc.save(path)

    result = parse_weekly_schedule(str(path), ["张三", "周一"])
    assert result["周一"] == ["张三"]


def test_parse_weekly_requires_standalone_day_heading(tmp_path):
    """星期标题必须单独成行；同一行写“周一 姓名”不应被当作标题"""
    from docx import Document

    path = tmp_path / "weekly_inline_day.docx"
    doc = Document()
    doc.add_paragraph("周一 张　爽 李四")
    doc.save(path)

    result = parse_weekly_schedule(str(path), ["张爽", "李四", "周一"])
    assert result["周一"] == []


def test_parse_weekly_standalone_day_then_fullwidth_space_name(tmp_path):
    """星期标题单独成行后，下一行姓名含全角空格时仍应正确提取"""
    from docx import Document

    path = tmp_path / "weekly_standalone_fullwidth.docx"
    doc = Document()
    doc.add_paragraph("周一")
    doc.add_paragraph("1-2节 张　爽 李四")
    doc.save(path)

    result = parse_weekly_schedule(str(path), ["张爽", "李四", "周一"])
    assert result["周一"] == ["张爽", "李四"]


def test_parse_weekly_name_containing_weekday_not_misread_as_heading(tmp_path):
    """姓名中包含“周一”等字样时，不应切换当前星期"""
    from docx import Document

    path = tmp_path / "weekly_name_contains_day.docx"
    doc = Document()
    doc.add_paragraph("周一")
    doc.add_paragraph("1-2节 周一二 李四")
    doc.save(path)

    result = parse_weekly_schedule(str(path), ["周一二", "李四"])
    assert result["周一"] == ["周一二", "李四"]


def test_count_actual_shifts_with_duplicates(make_weekly_docx):
    """同一人一周内多次排班，次数应累加"""
    schedule = {
        "周一": [["张三", "李四"]],
        "周二": [["张三"]],
        "周三": [["张三", "李四"]],
    }
    path = make_weekly_docx(schedule)
    total = ["张三", "李四", "王五"]
    weekly = parse_weekly_schedule(path, total)
    counter = count_actual_shifts(weekly, total)

    assert counter["张三"] == 3
    assert counter["李四"] == 2
    assert counter["王五"] == 0
