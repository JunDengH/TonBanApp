"""
word_generator 模块测试（需求3 第（三）项）。
覆盖：
- 首次生成（无上月）
- 上月多值/缺班继承
- 人员流动（离职丢弃 / 新进者按 0）
- 上月 Word 反向解析：空格单元格 / 斜线单元格 → 0
- 排序规则（多值降 / 缺班升 / 拼音升）
- 序号从 1 递增
- ≤0 单元格：文本为空 + tl2br 斜线注入
- 内置版式复刻（页面/表格/编号/斜线按黄金样本生成）
"""
from zipfile import ZipFile

import pytest
from lxml import etree
from docx import Document
from docx.oxml.ns import qn

from src.modules.word_generator import generate_monthly_word
from src.modules.word_parser import parse_previous_month_word


# ============================================================
# 读取辅助
# ============================================================
def _read_month_word(path):
    """返回生成后的月度Word中(表头之后)的所有数据行，解析为 list[dict]"""
    doc = Document(str(path))
    t = doc.tables[0]
    # 月统计Word：row0=标题，row1=表头，row2+=数据
    data = []
    for row in t.rows[2:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 6:
            continue
        data.append({
            "序号": cells[0],
            "姓名": cells[1],
            "总计班次": cells[2],
            "应值班次": cells[3],
            "多值班次": cells[4],  # 可能为空
            "缺班数量": cells[5],  # 可能为空
        })
    return data, t


def _is_auto_numbered(cell_obj):
    p = cell_obj._tc.find(qn("w:p"))
    if p is None:
        return False
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        return False
    numPr = pPr.find(qn("w:numPr"))
    return numPr is not None


def _has_tl2br(cell_obj):
    tcPr = cell_obj._tc.find(qn("w:tcPr"))
    if tcPr is None:
        return False
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        return False
    return borders.find(qn("w:tl2br")) is not None


XML_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _xml_qn(tag):
    prefix, name = tag.split(":")
    return f"{{{XML_NS[prefix]}}}{name}"


def _read_docx_xml(docx_path, part_name):
    with ZipFile(docx_path) as zf:
        return etree.fromstring(zf.read(part_name))


def _attr(el, name):
    return None if el is None else el.get(_xml_qn(name))


# ============================================================
# 1. 首次（无上月）
# ============================================================
def test_first_month_no_prev(tmp_path, make_weekly_excel):
    total = ["张三", "李四", "王五"]
    w1 = make_weekly_excel("w1.xlsx", total, {
        "周一": ["张三", "李四"],
        "周二": ["张三", "王五"],
    })
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, None, str(out))

    data, _ = _read_month_word(out)
    by_name = {r["姓名"]: r for r in data}
    # 张三 实际2 应值2 -> 总计2 应值2 多值=0空斜 缺班=0空斜
    assert by_name["张三"]["总计班次"] == "2"
    assert by_name["张三"]["应值班次"] == "2"
    assert by_name["张三"]["多值班次"] == ""
    assert by_name["张三"]["缺班数量"] == ""


def test_first_month_equals_excel_values(tmp_path, make_weekly_excel):
    """不传上月时，各字段数值应与需求2 Excel 口径一致（仅公式层面）"""
    total = ["张三"]
    w1 = make_weekly_excel("w1.xlsx", total, {
        "周一": ["张三", "张三", "张三"],  # 实际3 应值2
    })
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, None, str(out))
    data, _ = _read_month_word(out)
    r = data[0]
    assert r["总计班次"] == "3"
    assert r["应值班次"] == "3"
    # 多值=1 显示数字；缺班=-1 <=0 显示空
    assert r["多值班次"] == ""
    assert r["缺班数量"] == ""


# ============================================================
# 2. 跨月结转
# ============================================================
def test_inherit_over_from_prev(tmp_path, make_weekly_excel, make_previous_month_word):
    total = ["张三"]
    prev = make_previous_month_word([
        {"姓名": "张三", "多值班次": 5, "缺班数量": 0},
    ])
    w1 = make_weekly_excel("w1.xlsx", total, {
        "周一": ["张三", "张三"],  # 实际2 应值2
    })
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, prev, str(out))
    r = _read_month_word(out)[0][0]
    # 总计=2+5=7；应值=2+0=2；多值=5；缺班<=0 空
    assert r["总计班次"] == "7"
    assert r["应值班次"] == "2"
    assert r["多值班次"] == "5"
    assert r["缺班数量"] == ""


def test_inherit_absence_from_prev(tmp_path, make_weekly_excel, make_previous_month_word):
    total = ["李四"]
    prev = make_previous_month_word([
        {"姓名": "李四", "多值班次": 0, "缺班数量": 3},
    ])
    w1 = make_weekly_excel("w1.xlsx", total, {
        "周一": ["李四"],  # 实际1 应值2
    })
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, prev, str(out))
    r = _read_month_word(out)[0][0]
    # 总计=1；应值=1+3=4；多值<=0 空；缺班=3
    assert r["总计班次"] == "1"
    assert r["应值班次"] == "4"
    assert r["多值班次"] == ""
    assert r["缺班数量"] == "3"


# ============================================================
# 3. 人员流动
# ============================================================
def test_departed_dropped(tmp_path, make_weekly_excel, make_previous_month_word):
    """上月有'老赵'，本月总名单中无 → 不出现在结果里"""
    total = ["张三"]
    prev = make_previous_month_word([
        {"姓名": "张三", "多值班次": 2, "缺班数量": 0},
        {"姓名": "老赵", "多值班次": 9, "缺班数量": 0},
    ])
    w1 = make_weekly_excel("w1.xlsx", total, {"周一": ["张三"]})
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, prev, str(out))
    data, _ = _read_month_word(out)
    names = [r["姓名"] for r in data]
    assert "老赵" not in names
    assert "张三" in names


def test_new_joiner_zero(tmp_path, make_weekly_excel, make_previous_month_word):
    """本月新增的助理，上月多值/缺班按 0"""
    total = ["张三", "小王"]  # 小王是新进
    prev = make_previous_month_word([
        {"姓名": "张三", "多值班次": 2, "缺班数量": 0},
        # 小王不在上月数据里
    ])
    w1 = make_weekly_excel("w1.xlsx", total, {
        "周一": ["张三", "小王", "小王"],  # 小王实际2 应值2
    })
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, prev, str(out))
    by = {r["姓名"]: r for r in _read_month_word(out)[0]}
    # 小王：继承 0/0；总计2 应值2 多值/缺班都 <=0
    assert by["小王"]["总计班次"] == "2"
    assert by["小王"]["应值班次"] == "2"
    assert by["小王"]["多值班次"] == ""
    assert by["小王"]["缺班数量"] == ""


# ============================================================
# 4. 上月 Word 反向解析：斜线/空 → 0
# ============================================================
def test_slash_cell_parsed_as_zero(make_previous_month_word):
    prev = make_previous_month_word(
        [{"姓名": "张三", "多值班次": 999, "缺班数量": 999}],
        special={"张三": {"多值班次", "缺班数量"}},  # 两列都改为斜线
    )
    parsed = parse_previous_month_word(prev)
    assert parsed["张三"]["多值班次"] == 0
    assert parsed["张三"]["缺班数量"] == 0


def test_empty_cell_parsed_as_zero(tmp_path):
    """手工构造只写了'姓名'列、其余为空的Word"""
    from docx import Document as _D
    p = tmp_path / "empty_prev.docx"
    doc = _D()
    t = doc.add_table(rows=2, cols=6)
    for i, h in enumerate(["序号", "姓名", "总计班次", "应值班次", "多值班次", "缺班数量"]):
        t.rows[0].cells[i].text = h
    t.rows[1].cells[1].text = "张三"
    # 其余留空
    doc.save(p)
    parsed = parse_previous_month_word(str(p))
    assert parsed == {"张三": {"多值班次": 0, "缺班数量": 0}}


def test_invalid_header_raises(tmp_path):
    from docx import Document as _D
    p = tmp_path / "bad.docx"
    doc = _D()
    t = doc.add_table(rows=2, cols=3)
    t.rows[0].cells[0].text = "A"
    t.rows[0].cells[1].text = "B"
    t.rows[0].cells[2].text = "C"
    doc.save(p)
    with pytest.raises(ValueError):
        parse_previous_month_word(str(p))


# ============================================================
# 5. 排序
# ============================================================
def test_sort_order(tmp_path, make_weekly_excel):
    """
        构造：
            A多值: 实际3 应值3 -> 全0
            B正常: 实际2 应值2 -> 全0
            C缺班: 实际0 应值0 -> 全0
            D缺班: 实际1 应值1 -> 全0
        期望顺序：多值和缺班都相同，按拼音升序
    """
    total = ["A多值", "B正常", "C缺班", "D缺班"]
    w1 = make_weekly_excel("w1.xlsx", total, {
        "周一": ["A多值", "A多值", "A多值", "B正常", "B正常", "D缺班"],
    })
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, None, str(out))
    names = [r["姓名"] for r in _read_month_word(out)[0]]
    assert names == ["A多值", "B正常", "C缺班", "D缺班"]


def test_sort_tie_pinyin(tmp_path, make_weekly_excel):
    """多值相同、缺班相同时按姓名拼音升序"""
    total = ["周一二", "安小一"]
    # 两人都实际2 应值2 -> 多值0 缺班0 -> 按拼音：'a'<'z'，安小一在前
    w1 = make_weekly_excel("w1.xlsx", total, {
        "周一": ["周一二", "周一二", "安小一", "安小一"],
    })
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, None, str(out))
    names = [r["姓名"] for r in _read_month_word(out)[0]]
    assert names == ["安小一", "周一二"]


# ============================================================
# 6. 序号
# ============================================================
def test_serial_number_from_one(tmp_path, make_weekly_excel):
    total = [f"人{i}" for i in range(5)]
    w1 = make_weekly_excel("w1.xlsx", total, {"周一": ["人0"]})
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, None, str(out))
    _, t = _read_month_word(out)
    # 序号列改为Word自动编号：文本可为空，但必须带 numPr
    for ri in range(2, 2 + len(total)):
        assert _is_auto_numbered(t.rows[ri].cells[0])


# ============================================================
# 7. 斜线注入
# ============================================================
def test_diagonal_injected_for_non_positive(tmp_path, make_weekly_excel):
    """多值/缺班 <=0 的单元格：文本为空 + 存在 tl2br"""
    total = ["张三"]
    w1 = make_weekly_excel("w1.xlsx", total, {
        "周一": ["张三", "张三"],  # 实际2 应值2 -> 多值0 缺班0 双斜线
    })
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, None, str(out))
    _, t = _read_month_word(out)
    data_row = t.rows[2]
    # 第5列=多值班次，第6列=缺班数量
    over_cell = data_row.cells[4]
    absence_cell = data_row.cells[5]
    assert over_cell.text.strip() == ""
    assert absence_cell.text.strip() == ""
    assert _has_tl2br(over_cell)
    assert _has_tl2br(absence_cell)


def test_no_diagonal_when_positive(tmp_path, make_weekly_excel):
    """当值>0 时，不应保留斜线占位"""
    total = ["张三"]
    w1 = make_weekly_excel(
        "w1.xlsx",
        total,
        {"周一": ["张三"]},  # 应值1
        actual_schedule={"周一": ["张三", "张三", "张三"]},  # 实际3 -> 多值2
    )
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, None, str(out))
    _, t = _read_month_word(out)
    over_cell = t.rows[2].cells[4]
    assert over_cell.text.strip() == "2"
    assert not _has_tl2br(over_cell)


# ============================================================
# 8. 内置版式与标题
# ============================================================
def test_title_replaced(tmp_path, make_weekly_excel):
    total = ["张三"]
    w1 = make_weekly_excel("w1.xlsx", total, {"周一": ["张三"]})
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, None, str(out),
                          title_text="测试标题123")
    doc = Document(str(out))
    t = doc.tables[0]
    assert "测试标题123" in t.rows[0].cells[0].text


def test_header_row_preserved(tmp_path, make_weekly_excel):
    """表头行不能被误删/改写"""
    total = ["张三"]
    w1 = make_weekly_excel("w1.xlsx", total, {"周一": ["张三"]})
    out = tmp_path / "m.docx"
    generate_monthly_word([w1], total, None, str(out))
    doc = Document(str(out))
    t = doc.tables[0]
    headers = [c.text.strip() for c in t.rows[1].cells]
    assert headers == ["序号", "姓名", "总计班次", "应值班次", "多值班次", "缺班数量"]


def test_forced_title_format(tmp_path, make_weekly_excel):
    """标题可被外部强制为指定格式"""
    total = ["张三"]
    w1 = make_weekly_excel("w1.xlsx", total, {"周一": ["张三"]})
    out = tmp_path / "m.docx"
    title = "2026春季学期1-4周助理值班统计"
    generate_monthly_word([w1], total, None, str(out), title_text=title)
    doc = Document(str(out))
    t = doc.tables[0]
    assert t.rows[0].cells[0].text == title


def test_monthly_overtime_added_to_total(tmp_path, make_weekly_excel):
    """加班补录在月统计阶段叠加到总计班次"""
    total = ["张三", "李四"]
    w1 = make_weekly_excel("w1.xlsx", total, {
        "周一": ["张三"],  # 原始实际：张三1，李四0
    })
    out = tmp_path / "m_overtime.docx"
    generate_monthly_word(
        [w1],
        total,
        None,
        str(out),
        overtime_shifts={"张三": 2, "李四": 1},
    )

    data, _ = _read_month_word(out)
    by = {r["姓名"]: r for r in data}
    assert by["张三"]["总计班次"] == "3"
    assert by["李四"]["总计班次"] == "1"


def test_generated_word_embeds_golden_layout_contract(tmp_path, make_weekly_excel):
    """月统计Word不依赖外部模板，生成件应内置5-8周黄金样本的关键版式。"""
    total = ["张三", "李四"]
    w1 = make_weekly_excel("w1.xlsx", total, {
        "周一": ["张三"],
        "周二": ["李四", "李四"],
    })
    out = tmp_path / "layout_contract.docx"
    generate_monthly_word(
        [w1],
        total,
        None,
        str(out),
        title_text="2026春季学期5-8周助理值班统计",
    )

    doc_xml = _read_docx_xml(out, "word/document.xml")
    sect = doc_xml.xpath(".//w:sectPr", namespaces=XML_NS)[-1]
    pg_sz = sect.find(_xml_qn("w:pgSz"))
    pg_mar = sect.find(_xml_qn("w:pgMar"))
    assert {k: _attr(pg_sz, f"w:{k}") for k in ("w", "h")} == {
        "w": "11906",
        "h": "16838",
    }
    assert {k: _attr(pg_mar, f"w:{k}") for k in ("top", "right", "bottom", "left")} == {
        "top": "1440",
        "right": "1800",
        "bottom": "1440",
        "left": "1800",
    }

    table = doc_xml.xpath(".//w:tbl", namespaces=XML_NS)[0]
    grid = table.find(_xml_qn("w:tblGrid"))
    assert [_attr(col, "w:w") for col in grid.findall(_xml_qn("w:gridCol"))] == [
        "1086",
        "1663",
        "1708",
        "1715",
        "1774",
        "1778",
    ]

    rows = table.findall(_xml_qn("w:tr"))
    heights = []
    for row in rows:
        tr_pr = row.find(_xml_qn("w:trPr"))
        height = tr_pr.find(_xml_qn("w:trHeight"))
        heights.append((_attr(height, "w:val"), _attr(height, "w:hRule")))
    assert heights[0] == ("983", "atLeast")
    assert heights[1] == ("519", "atLeast")
    assert heights[2:] == [("561", "atLeast"), ("561", "atLeast")]

    numbering_xml = _read_docx_xml(out, "word/numbering.xml")
    num = numbering_xml.find(_xml_qn("w:num"))
    assert _attr(num, "w:numId") == "1"
    assert _attr(num.find(_xml_qn("w:abstractNumId")), "w:val") == "0"
    lvl0 = numbering_xml.find(_xml_qn("w:abstractNum")).find(_xml_qn("w:lvl"))
    assert _attr(lvl0, "w:ilvl") == "0"
    assert _attr(lvl0.find(_xml_qn("w:numFmt")), "w:val") == "decimal"
    assert _attr(lvl0.find(_xml_qn("w:lvlText")), "w:val") == "%1"
    assert _attr(lvl0.find(_xml_qn("w:lvlJc")), "w:val") == "center"

    first_no_cell = rows[2].findall(_xml_qn("w:tc"))[0]
    num_pr = first_no_cell.find(".//w:numPr", namespaces=XML_NS)
    assert _attr(num_pr.find(_xml_qn("w:ilvl")), "w:val") == "0"
    assert _attr(num_pr.find(_xml_qn("w:numId")), "w:val") == "1"
