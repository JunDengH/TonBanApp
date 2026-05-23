import pytest
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@pytest.fixture
def make_total_names_docx(tmp_path):
    """生成一个总名单Word"""
    def _factory(names):
        path = tmp_path / "total.docx"
        doc = Document()
        doc.add_paragraph(" ".join(names))
        doc.save(path)
        return str(path)
    return _factory


@pytest.fixture
def make_weekly_docx(tmp_path):
    """
    生成周统计Word
    schedule: {"周一": [["张三","李四"], ["王五","赵六"]], "周二": [...]}
    外层列表=当天的多个时段，内层=每个时段的排班人名
    """
    def _factory(schedule):
        path = tmp_path / "weekly.docx"
        doc = Document()
        for day, sessions in schedule.items():
            doc.add_paragraph(day)
            for i, names in enumerate(sessions):
                line = f"{1+i*2}-{2+i*2}节 " + " ".join(names)
                doc.add_paragraph(line)
        doc.save(path)
        return str(path)
    return _factory


# ============================================================
# 月度 Word fixture（需求3 反向解析测试用）
# ============================================================
def _add_tl2br(cell):
    """向单元格注入左上→右下斜线"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    tl2br = OxmlElement("w:tl2br")
    tl2br.set(qn("w:val"), "single")
    tl2br.set(qn("w:sz"), "4")
    tl2br.set(qn("w:color"), "000000")
    borders.append(tl2br)


@pytest.fixture
def make_previous_month_word(tmp_path):
    """
    构造一份"上月月度统计"Word，用于反向解析测试。

    records: list[dict]，每条形如：
        {"姓名": "张三", "多值班次": 2, "缺班数量": 0}
    special: dict[str, set[str]]，可选，指定哪些人的哪些列用"斜线/空"代替数字：
        {"李四": {"多值班次", "缺班数量"}, ...}
    filename: 可选文件名，默认 prev_month.docx
    """
    def _factory(records, special=None, filename="prev_month.docx"):
        special = special or {}
        path = tmp_path / filename
        doc = Document()
        doc.add_paragraph("助理值班统计")
        table = doc.add_table(rows=1 + len(records), cols=6)
        # 表头
        headers = ["序号", "姓名", "总计班次", "应值班次", "多值班次", "缺班数量"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        # 数据
        for ri, rec in enumerate(records, start=1):
            cells = table.rows[ri].cells
            cells[0].text = str(ri)
            cells[1].text = rec["姓名"]
            cells[2].text = str(rec.get("总计班次", 0))
            cells[3].text = str(rec.get("应值班次", 0))

            over_special = "多值班次" in special.get(rec["姓名"], set())
            absence_special = "缺班数量" in special.get(rec["姓名"], set())

            # 多值班次
            if over_special:
                cells[4].text = ""
                _add_tl2br(cells[4])
            else:
                cells[4].text = str(rec.get("多值班次", 0))

            # 缺班数量
            if absence_special:
                cells[5].text = ""
                _add_tl2br(cells[5])
            else:
                cells[5].text = str(rec.get("缺班数量", 0))

        doc.save(path)
        return str(path)
    return _factory


@pytest.fixture
def make_weekly_excel(tmp_path):
    """
    直接用 generate_weekly_excel + 伪造 parse_weekly_schedule 生成周 Excel。
    供月度 Word 测试快速构造输入。
    schedule: {"周一": ["张三", ...], ...}
    """
    from unittest.mock import patch
    from src.modules.excel_generator import generate_weekly_excel

    def _factory(filename, total_names, schedule, holidays=None, actual_schedule=None):
        output = tmp_path / filename
        actual_schedule = actual_schedule if actual_schedule is not None else schedule
        with patch(
            "src.modules.excel_generator.parse_weekly_schedule",
            side_effect=[schedule, actual_schedule],
        ):
            generate_weekly_excel(
                "fake.docx", total_names, holidays or [], str(output),
                actual_word_path="fake.docx"
            )
        return str(output)

    return _factory