# src/modules/word_parser.py
"""
Word解析模块：
1. 解析总名单Word（按空格分隔提取姓名）
2. 解析周统计Word（按周一~周日提取每日值班名单）
"""
import re
from docx import Document
from docx.oxml.ns import qn

from src.utils.helpers import chars_homophone, chars_sound_similar, edit_distance


# 中文姓名匹配：2-4个汉字
NAME_REGEX = re.compile(r"[\u4e00-\u9fa5]{2,4}")

# 月度Word表头列名（反向解析时用于匹配列索引）
MONTHLY_HEADER_REQUIRED = ("姓名", "多值班次", "缺班数量")
# 识别星期标题：周一、星期一、周1 等。
# 约定：星期标题必须在 Word 中单独成行，避免姓名中包含“周一”等词时被误判。
WEEKDAY_PATTERNS = {
    "周一": ["周一", "星期一", "礼拜一", "周1"],
    "周二": ["周二", "星期二", "礼拜二", "周2"],
    "周三": ["周三", "星期三", "礼拜三", "周3"],
    "周四": ["周四", "星期四", "礼拜四", "周4"],
    "周五": ["周五", "星期五", "礼拜五", "周5"],
    "周六": ["周六", "星期六", "礼拜六", "周6"],
    "周日": ["周日", "星期日", "星期天", "礼拜日", "礼拜天", "周7"],
}


def _compact_text(value: str) -> str:
    """Remove all whitespace from text for compact name matching."""
    return re.sub(r"\s+", "", str(value or ""))


_WEEKDAY_ANNOTATION_RE = re.compile(
    r"^[（(].*[\)）]$"
    r"|"
    r"^[\d.月日号/\-:：~～\s]*$"
)


def _is_weekday_annotation(text: str) -> bool:
    """Check if text looks like a date/time annotation after a weekday alias."""
    if not text:
        return True
    return bool(_WEEKDAY_ANNOTATION_RE.match(text))


def _match_roster_names(text: str, total_names: list, corrections: dict | None = None) -> list:
    """Match roster names in text by compacting whitespace, then greedy prefix matching.

    Handles names with internal spaces (e.g. "冯  泽" -> "冯泽") consistently with
    the PDF parser.  Roster names are matched longest-first to avoid prefix collisions.

    When *corrections* is provided, each typo key is treated as an additional alias
    for its mapped roster name, so "张灏琛" in the text would be counted as "张颢琛"
    without modifying the source file.
    """
    compact = _compact_text(text)
    candidates = list(
        (_compact_text(name), name) for name in total_names if _compact_text(name)
    )
    if corrections:
        for typo, correct in corrections.items():
            compact_typo = _compact_text(typo)
            if compact_typo and compact_typo != _compact_text(correct):
                candidates.append((compact_typo, correct))
    candidates.sort(key=lambda item: len(item[0]), reverse=True)
    matches = []
    index = 0
    while index < len(compact):
        matched = None
        for normalized, original in candidates:
            if compact.startswith(normalized, index):
                matched = (normalized, original)
                break
        if matched is None:
            index += 1
            continue
        matches.append(matched[1])
        index += len(matched[0])
    return matches


def _unknown_candidates(text: str, total_names: list, corrections: dict | None = None) -> list:
    """Find 2-4 char Chinese sequences in text that are not roster names or weekday aliases.

    When *corrections* is provided, the typo keys are also removed from the Chinese
    text so that confirmed-correction names don't show up as "unknown".
    """
    removal_set = {_compact_text(name) for name in total_names if _compact_text(name)}
    if corrections:
        removal_set |= {_compact_text(typo) for typo in corrections}
    normalized_names = sorted(removal_set, key=len, reverse=True)
    weekday_alias_set = {
        _compact_text(alias)
        for aliases in WEEKDAY_PATTERNS.values()
        for alias in aliases
        if _compact_text(alias)
    }
    candidates = []
    for line in str(text or "").splitlines():
        chinese_only = "".join(re.findall(r"[\u4e00-\u9fa5]", line))
        for name in normalized_names:
            chinese_only = chinese_only.replace(name, " ")
        for token in NAME_REGEX.findall(chinese_only):
            if token not in weekday_alias_set:
                candidates.append(token)
    return candidates


def _normalize_name_line(line: str) -> str:
    """
    规范化姓名行文本：
    - 仅移除中文字符之间的全角空格（\u3000），修复“张　爽”这类被拆分问题；
    - 不移除普通空格，避免把“张三 李四”错误拼接成“张三李四”。
    """
    if not line:
        return ""
    return re.sub(r"(?<=[\u4e00-\u9fa5])\u3000(?=[\u4e00-\u9fa5])", "", line)


def _extract_names_from_line(line: str) -> list:
    """从单行文本中提取姓名（2-4 个汉字），兼容姓名内全角空格。"""
    normalized = _normalize_name_line(line)
    names = []
    for token in re.split(r"\s+", normalized.strip()):
        if not token:
            continue
        names.extend(NAME_REGEX.findall(token))
    return names


def _strip_weekday_marker(line: str, weekday_key: str) -> str:
    """
    去掉单独成行的周标题别名，避免“周几/星期几”等被误识别为姓名。
    """
    if not line:
        return ""
    aliases = WEEKDAY_PATTERNS.get(weekday_key, [])
    cleaned = line
    for alias in aliases:
        cleaned = cleaned.replace(alias, " ")
    return cleaned


def _iter_all_text(doc: Document):
    """遍历文档：正文段落 + 所有表格单元格，返回文本行列表"""
    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if text:
                        lines.append(text)
    return lines


def parse_total_name_list(docx_path: str) -> list:
    """
    解析总名单：提取所有中文姓名
    要求：Word中姓名以空格分隔
    """
    doc = Document(docx_path)
    lines = _iter_all_text(doc)

    names = []
    seen = set()
    for line in lines:
        for name in _extract_names_from_line(line):
            if name not in seen:
                seen.add(name)
                names.append(name)
    return names


def _match_weekday(text: str) -> str:
    """判断一段文本是否是某一天的标题行；星期标题必须单独成行。

    支持带日期注解的星期标题，如 "周一（6.29）"、"周二(7.30)"、"周三 6.29"，
    但 "周一 张三"（姓名跟在后面）不视为星期标题。
    """
    normalized = (text or "").strip()
    for key, aliases in WEEKDAY_PATTERNS.items():
        for alias in aliases:
            if normalized == alias:
                return key
            if normalized.startswith(alias):
                rest = normalized[len(alias):].strip()
                if _is_weekday_annotation(rest):
                    return key
    return None


def parse_weekly_schedule(docx_path: str, total_names: list, corrections: dict | None = None) -> dict:
    """
    解析周统计Word，返回：
    {
        "周一": ["张三", "李四", "张三", ...],   # 保留重复，表示多次排班
        "周二": [...],
        ...
    }
    说明：
        - 按出现顺序遇到单独成行的周X标题，其后（直到下一个周X标题）的人名都归入该天；
        - 若文档通过表格呈现，会将所有单元格文本按顺序纳入；
        - 只保留出现在总名单中的姓名，自动过滤"1-2节"等课节干扰。
        - corrections: 暂时纠错映射 {"张灏琛": "张颢琛"}，将错字当作正确姓名计数。
    """
    doc = Document(docx_path)
    lines = _iter_all_text(doc)

    result = {k: [] for k in WEEKDAY_PATTERNS.keys()}
    current_day = None

    for line in lines:
        weekday = _match_weekday(line)
        if weekday is not None:
            current_day = weekday
            # 星期标题必须单独成行；标题行本身通常不包含姓名。
        if current_day is None:
            continue

        parse_line = line
        if weekday is not None:
            parse_line = _strip_weekday_marker(line, weekday)

        # 提取中文姓名并按总名单过滤（压缩空白后匹配，兼容 "冯  泽" 这类带内部空格的姓名）
        result[current_day].extend(_match_roster_names(parse_line, total_names, corrections))

    return result


def count_actual_shifts(weekly_schedule: dict, total_names: list) -> dict:
    """统计每位助理在一周内的实际出现次数"""
    counter = {n: 0 for n in total_names}
    for day, names in weekly_schedule.items():
        for n in names:
            if n in counter:
                counter[n] += 1
    return counter


def scan_weekly_word(docx_path: str, total_names: list, corrections: dict | None = None) -> dict:
    """
    扫描周排班/实际 Word，用于生成前的合理性校验。

    返回:
        {
            "matched_count": int,        # 落在总名单中的姓名出现总次数
            "unknown_names": [str, ...], # 出现在日分区里、形如姓名（2-4 汉字）
                                          # 但不在总名单、且非星期别名的 token（去重、保序）
        }

    说明:
        - 与 parse_weekly_schedule 走同一套日分区解析逻辑；
        - unknown_names 仅作软警告参考，可能混入"签到表"等说明文字，调用方应注明。
        - corrections: 暂时纠错映射，已纠错的姓名计入 matched_count 且不出现在 unknown_names。
    """
    doc = Document(docx_path)
    lines = _iter_all_text(doc)

    weekday_aliases = {alias for aliases in WEEKDAY_PATTERNS.values() for alias in aliases}
    matched_count = 0
    unknown_names = []
    seen_unknown = set()
    current_day = None

    for line in lines:
        weekday = _match_weekday(line)
        if weekday is not None:
            current_day = weekday
        if current_day is None:
            continue

        parse_line = line
        if weekday is not None:
            parse_line = _strip_weekday_marker(line, weekday)

        matched_count += len(_match_roster_names(parse_line, total_names, corrections))
        for candidate in _unknown_candidates(parse_line, total_names, corrections):
            if candidate not in weekday_aliases and candidate not in seen_unknown:
                seen_unknown.add(candidate)
                unknown_names.append(candidate)

    return {"matched_count": matched_count, "unknown_names": unknown_names}


# 疑似档位排序：strong（高度疑似）> medium（疑似）> weak（形近待确认）
TYPO_LEVEL_ORDER = {"strong": 3, "medium": 2, "weak": 1}


def _classify_typo(unknown: str, total: str):
    """
    判断未知名 unknown 相对总名单姓名 total 的疑似打错字档位，命中返回档位字符串，否则 None。

    档位规则（详见 plan，已与用户确认）：
      - strong：长度相同、逐字读音都能匹配（多音字任一读音相交）但用字不同
                → 几乎肯定同音打错字（不受编辑距离限制，可覆盖多字皆错的双同音）。
      - medium：字符编辑距离 == 1，且
                · 漏字 / 多字（长度差 1，其余字相同）；或
                · 替换且被替换的字读音相近（同音或模糊音命中）。
      - weak：  字符编辑距离 == 1 的替换，但读音对不上（纯形近）。
                为避免遗漏，两字姓名同样进入此档（宁可多报不漏）。
    """
    if unknown == total:
        return None

    # strong：整名读音一致（用字不同）。
    if len(unknown) == len(total) and all(
        chars_homophone(cu, cm) for cu, cm in zip(unknown, total)
    ):
        return "strong"

    if edit_distance(unknown, total) != 1:
        return None

    # 漏字 / 多字：其余字完全相同，天然高可信。
    if len(unknown) != len(total):
        return "medium"

    # 替换：长度相同且编辑距离 1，恰有一个位置不同。
    cu, cm = next((u, t) for u, t in zip(unknown, total) if u != t)
    if chars_sound_similar(cu, cm):
        return "medium"
    return "weak"


def find_typo_suspects(unknown_names: list, total_names: list) -> list:
    """
    在“不在总名单”的疑似姓名里，找出与总名单某助理疑似打错字的项，并分档。

    参数:
        unknown_names: scan_weekly_word 返回的 unknown_names（保序、去重）
        total_names:   总名单姓名列表

    返回:
        [{"name": 未知名, "candidates": [总名单姓名, ...], "level": "strong"|"medium"|"weak"}, ...]
        - 仅含命中至少一个候选的未知名；顺序与入参一致；
        - 同一未知名命中多档时只保留最高档，candidates 按总名单顺序保留该档全部命中。
    """
    suspects = []
    for unknown in unknown_names:
        best_level = None
        candidates = []
        for total in total_names:
            level = _classify_typo(unknown, total)
            if level is None:
                continue
            if best_level is None or TYPO_LEVEL_ORDER[level] > TYPO_LEVEL_ORDER[best_level]:
                best_level = level
                candidates = [total]
            elif level == best_level:
                candidates.append(total)
        if best_level is not None:
            suspects.append({"name": unknown, "candidates": candidates, "level": best_level})
    return suspects


# ============================================================
# 上月月度 Word 反向解析（需求3 跨月结转使用）
# ============================================================
def _cell_has_diagonal(cell) -> bool:
    """判断单元格是否含左上→右下斜线（w:tl2br）"""
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        return False
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        return False
    return borders.find(qn("w:tl2br")) is not None


def _safe_int(text: str) -> int:
    """将单元格文本稳健转为整数；非数字一律按 0 处理"""
    if text is None:
        return 0
    s = str(text).strip()
    if not s:
        return 0
    # 去掉可能的空白与全角符号
    try:
        return int(float(s))
    except (ValueError, TypeError):
        return 0


def parse_previous_month_word(docx_path: str) -> dict:
    """
    解析上月月度统计 Word，返回每位助理的"多值班次""缺班数量"。

    反向解析规则（需求3 第4条 防篡改与占位）：
      - 单元格含 w:tl2br 斜线 -> 视为 0
      - 单元格文本为空         -> 视为 0
      - 单元格文本非整数       -> 视为 0
      - 单元格文本为负数       -> 视为 0（多值班次/缺班数量语义上非负，
                                 合法月报只会出现非负数或斜线占位）

    返回:
        {姓名: {"多值班次": int, "缺班数量": int}, ...}

    异常:
        ValueError: 文档中不存在含"姓名""多值班次""缺班数量"三列的表格时抛出
    """
    doc = Document(docx_path)
    if not doc.tables:
        raise ValueError("上月Word未发现任何表格，无法反向解析。")

    # 找到首个同时含 3 个必需列的表格，并定位其表头行
    target_table = None
    header_row_idx = -1
    col_map = {}
    for table in doc.tables:
        for ri, row in enumerate(table.rows):
            cells_text = [c.text.strip() for c in row.cells]
            if all(h in cells_text for h in MONTHLY_HEADER_REQUIRED):
                # 记录列索引（取第一次出现的索引）
                col_map = {h: cells_text.index(h) for h in MONTHLY_HEADER_REQUIRED}
                target_table = table
                header_row_idx = ri
                break
        if target_table is not None:
            break

    if target_table is None:
        raise ValueError(
            "上月Word表头不符合规范，需要至少包含【姓名 / 多值班次 / 缺班数量】三列。"
        )

    name_col = col_map["姓名"]
    over_col = col_map["多值班次"]
    absence_col = col_map["缺班数量"]

    result = {}
    for row in target_table.rows[header_row_idx + 1:]:
        cells = row.cells
        if len(cells) <= max(name_col, over_col, absence_col):
            continue
        name = cells[name_col].text.strip()
        if not name:
            continue
        # 过滤非人名（避免把标题行等也收进来）
        if not NAME_REGEX.fullmatch(name):
            continue

        over_cell = cells[over_col]
        absence_cell = cells[absence_col]

        over_val = 0 if _cell_has_diagonal(over_cell) else max(0, _safe_int(over_cell.text))
        absence_val = 0 if _cell_has_diagonal(absence_cell) else max(0, _safe_int(absence_cell.text))

        result[name] = {
            "多值班次": over_val,
            "缺班数量": absence_val,
        }
    return result
