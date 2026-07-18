"""月统计加班补录、罚班记录 Word 导入导出。"""
from __future__ import annotations

import datetime
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt


BASE_HEADERS = ["序号", "姓名"]
OVERTIME_SHIFT_HEADER = "加班班次"
PENALTY_COUNT_HEADER = "罚班次数"


def export_overtime_word(records: dict[str, int], output_dir: str | Path) -> str:
    """将月统计加班补录记录导出为 Word 文件。"""
    return _export_adjustment_word(
        records=records,
        output_dir=output_dir,
        title="月统计加班补录",
        value_header=OVERTIME_SHIFT_HEADER,
        filename_prefix="月统计加班补录",
        empty_message="当前没有可导出的加班补录记录。",
    )


def export_penalty_word(records: dict[str, int], output_dir: str | Path) -> str:
    """将月统计罚班记录导出为 Word 文件。"""
    return _export_adjustment_word(
        records=records,
        output_dir=output_dir,
        title="月统计罚班记录",
        value_header=PENALTY_COUNT_HEADER,
        filename_prefix="月统计罚班记录",
        empty_message="当前没有可导出的罚班记录。",
    )


def _export_adjustment_word(
    records: dict[str, int],
    output_dir: str | Path,
    title: str,
    value_header: str,
    filename_prefix: str,
    empty_message: str,
) -> str:
    clean_records = _normalize_records(records)
    if not clean_records:
        raise ValueError(empty_message)

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{filename_prefix}_{datetime.datetime.now():%Y%m%d_%H%M%S}.docx"

    doc = Document()
    title_paragraph = doc.add_paragraph(title)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_paragraph.runs:
        title_paragraph.runs[0].font.name = "宋体"
        title_paragraph.runs[0].font.size = Pt(16)
        title_paragraph.runs[0].font.bold = True

    headers = [*BASE_HEADERS, value_header]
    table = doc.add_table(rows=1 + len(clean_records), cols=len(headers))
    table.style = "Table Grid"
    for col_idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[col_idx], header)

    for row_idx, (name, value) in enumerate(clean_records.items(), start=1):
        cells = table.rows[row_idx].cells
        _set_cell_text(cells[0], str(row_idx))
        _set_cell_text(cells[1], name)
        _set_cell_text(cells[2], str(value))

    doc.save(str(output_path))
    return str(output_path)


def import_overtime_word(docx_path: str | Path, total_names: list[str] | None = None) -> dict[str, int]:
    """从 Word 文件导入月统计加班补录记录。"""
    return _import_adjustment_word(
        docx_path=docx_path,
        total_names=total_names,
        value_header=OVERTIME_SHIFT_HEADER,
        record_label="加班班次",
        empty_message="导入失败，Word 中没有有效的加班补录记录。",
    )


def import_penalty_word(docx_path: str | Path, total_names: list[str] | None = None) -> dict[str, int]:
    """从 Word 文件导入月统计罚班记录。"""
    return _import_adjustment_word(
        docx_path=docx_path,
        total_names=total_names,
        value_header=PENALTY_COUNT_HEADER,
        record_label="罚班次数",
        empty_message="导入失败，Word 中没有有效的罚班记录。",
    )


def _import_adjustment_word(
    docx_path: str | Path,
    total_names: list[str] | None,
    value_header: str,
    record_label: str,
    empty_message: str,
) -> dict[str, int]:
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"文件不存在: {docx_path}")

    doc = Document(str(docx_path))
    table, col_map = _find_adjustment_table(doc, value_header)
    valid_names = set(total_names or [])
    imported: dict[str, int] = {}
    invalid_names = []
    invalid_values = []

    header_row_idx = col_map.pop("_row")
    for row in table.rows[header_row_idx + 1:]:
        cells = row.cells
        name = cells[col_map["姓名"]].text.strip()
        value_text = cells[col_map[value_header]].text.strip()
        if not name and not value_text:
            continue
        if not name:
            invalid_names.append("空姓名")
            continue
        if valid_names and name not in valid_names:
            invalid_names.append(name)
            continue

        value = _parse_positive_int(value_text)
        if value is None:
            invalid_values.append(f"{name}: {value_text or '空'}")
            continue
        imported[name] = imported.get(name, 0) + value

    if invalid_names:
        names = "、".join(dict.fromkeys(invalid_names))
        raise ValueError(f"导入失败，以下姓名不在总名单中或为空: {names}")
    if invalid_values:
        details = "；".join(invalid_values)
        raise ValueError(f"导入失败，{record_label}必须为正整数: {details}")
    if not imported:
        raise ValueError(empty_message)
    return imported


def _find_adjustment_table(doc: Document, value_header: str):
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if all(header in cells for header in ("姓名", value_header)):
                return table, {
                    "_row": row_idx,
                    "姓名": cells.index("姓名"),
                    value_header: cells.index(value_header),
                }
    raise ValueError(f"导入失败，未找到包含“姓名”和“{value_header}”表头的 Word 表格。")


def _normalize_records(records: dict[str, int]) -> dict[str, int]:
    clean: dict[str, int] = {}
    for name, shifts in records.items():
        name = str(name).strip()
        parsed = _parse_positive_int(str(shifts))
        if name and parsed is not None:
            clean[name] = clean.get(name, 0) + parsed
    return clean


def _parse_positive_int(value: str) -> int | None:
    text = str(value or "").strip()
    if not text:
        return None
    if not re.fullmatch(r"\d+", text):
        return None
    number = int(text)
    return number if number > 0 else None


def _set_cell_text(cell, text: str) -> None:
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "宋体"
            run.font.size = Pt(11)
